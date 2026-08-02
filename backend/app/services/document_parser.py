from pathlib import Path
import csv
from html.parser import HTMLParser
from io import StringIO
import json
from typing import Any, Protocol

from backend.app.core.config import DOCUMENT_LOADER_ORDER
from backend.app.services.document_types import LoaderResult, PARSER_VERSION, ParsedBlock, ParsedDocument


class DocumentParseError(ValueError):
    pass


class DocumentLoader(Protocol):
    name: str

    def supports(self, file_path: Path) -> bool: ...

    def load(self, file_path: Path) -> LoaderResult: ...


def parse_document_text(file_path: Path) -> str:
    """Compatibility wrapper for callers that still need a plain text string."""

    return parse_document(file_path).text


def parse_document(
    file_path: Path,
    loader_name: str | None = None,
    source_name: str | None = None,
) -> ParsedDocument:
    """Extract structured text blocks through the configured loader adapter chain."""

    errors: list[str] = []
    for loader in _candidate_loaders(file_path, loader_name):
        try:
            result = loader.load(file_path)
            blocks = _ensure_blocks(result.blocks)
            if source_name:
                _apply_source_name(blocks, result, source_name)
            return ParsedDocument(
                text=_join_blocks(blocks),
                blocks=blocks,
                metadata={
                    **result.metadata,
                    "source_format": file_path.suffix.lower().lstrip("."),
                    "parser_version": PARSER_VERSION,
                    "loader_name": result.loader_name,
                    "block_count": len(blocks),
                    "source_filename": source_name or file_path.name,
                },
            )
        except DocumentParseError as exc:
            errors.append(f"{loader.name}: {exc}")
            if loader_name is not None:
                break

    if errors:
        raise DocumentParseError("；".join(errors))
    raise DocumentParseError("暂不支持该文档格式")


def _apply_source_name(blocks: list[ParsedBlock], result: LoaderResult, source_name: str) -> None:
    source_path = Path(source_name)
    stem = source_path.stem
    source_title = stem.split("_", maxsplit=1)[-1] if "_" in stem else stem
    result.metadata["source_title"] = source_title
    result.metadata["source_filename"] = source_name


def parse_document_with_loader(file_path: Path, loader_name: str) -> ParsedDocument:
    """Run a specific loader by name for loader evaluation."""

    return parse_document(file_path=file_path, loader_name=loader_name)


def available_loader_names(file_path: Path) -> list[str]:
    return [loader.name for loader in _candidate_loaders(file_path)]


class TextDocumentLoader:
    name = "text"

    def supports(self, file_path: Path) -> bool:
        return file_path.suffix.lower() == ".txt"

    def load(self, file_path: Path) -> LoaderResult:
        text = _read_utf8_text(file_path)
        normalized = _normalize_text(text)
        blocks = [ParsedBlock(text=normalized, block_type="paragraph", order_index=1)] if normalized else []
        return LoaderResult(blocks=blocks, loader_name=self.name)


class MarkdownDocumentLoader:
    name = "markdown"

    def supports(self, file_path: Path) -> bool:
        return file_path.suffix.lower() == ".md"

    def load(self, file_path: Path) -> LoaderResult:
        text = _read_utf8_text(file_path)
        return LoaderResult(blocks=_markdown_to_blocks(text), loader_name=self.name)


class CsvDocumentLoader:
    name = "spreadsheet-csv"

    def supports(self, file_path: Path) -> bool:
        return file_path.suffix.lower() == ".csv"

    def load(self, file_path: Path) -> LoaderResult:
        text = _read_utf8_text(file_path)
        try:
            rows = list(csv.reader(StringIO(text)))
        except csv.Error as exc:
            raise DocumentParseError("CSV 表格解析失败") from exc
        blocks: list[ParsedBlock] = []
        for row in rows:
            line = "，".join(cell for cell in (_normalize_text(cell) for cell in row) if cell)
            if not line:
                continue
            blocks.append(
                ParsedBlock(
                    text=line,
                    block_type="paragraph",
                    order_index=len(blocks) + 1,
                )
            )
        if not blocks:
            raise DocumentParseError("CSV 没有可解析数据")
        return LoaderResult(
            blocks=blocks,
            loader_name=self.name,
            metadata={"source_format": "csv", "row_count": len(blocks)},
        )


class JsonlDocumentLoader:
    name = "jsonl"

    def supports(self, file_path: Path) -> bool:
        return file_path.suffix.lower() == ".jsonl"

    def load(self, file_path: Path) -> LoaderResult:
        text = _read_utf8_text(file_path)
        blocks: list[ParsedBlock] = []
        for line_number, line in enumerate(text.splitlines(), start=1):
            if not line.strip():
                continue
            try:
                record = json.loads(line)
            except json.JSONDecodeError as exc:
                raise DocumentParseError(f"JSONL 第 {line_number} 行不是有效 JSON") from exc
            if not isinstance(record, dict):
                raise DocumentParseError(f"JSONL 第 {line_number} 行必须是对象")
            keys = {str(key).lower() for key in record}
            if "question" in keys and keys & {"answer", "answer_text", "evidence", "option_a"}:
                raise DocumentParseError("检测到 QA/答案数据；评测数据禁止作为知识库文档入库")
            body_key = next((key for key in ("text", "content", "body") if record.get(key)), None)
            if body_key is None:
                raise DocumentParseError(
                    f"JSONL 第 {line_number} 行缺少 text、content 或 body 字段"
                )
            body = _normalize_text(str(record[body_key]))
            if not body:
                continue
            title = _normalize_text(str(record.get("title") or "")) or None
            metadata = {
                str(key): value
                for key, value in record.items()
                if key not in {"text", "content", "body"} and _json_metadata_value(value)
            }
            metadata["jsonl_record_index"] = line_number
            blocks.append(
                ParsedBlock(
                    text=body,
                    block_type="paragraph",
                    order_index=len(blocks) + 1,
                    section_title=title,
                    metadata=metadata,
                )
            )
        return LoaderResult(
            blocks=blocks,
            loader_name=self.name,
            metadata={"record_count": len(blocks), "source_format": "jsonl"},
        )


class HtmlDocumentLoader:
    name = "html"

    def supports(self, file_path: Path) -> bool:
        return file_path.suffix.lower() in {".html", ".htm"}

    def load(self, file_path: Path) -> LoaderResult:
        parser = _StructuredHtmlParser()
        try:
            parser.feed(_read_utf8_text(file_path))
            parser.close()
        except Exception as exc:
            raise DocumentParseError("HTML 正文解析失败") from exc
        return LoaderResult(
            blocks=parser.blocks,
            loader_name=self.name,
            metadata={"html_title": parser.title, "source_format": "html"},
        )


class DocxDocumentLoader:
    name = "python-docx"

    def supports(self, file_path: Path) -> bool:
        return file_path.suffix.lower() == ".docx"

    def load(self, file_path: Path) -> LoaderResult:
        try:
            from docx import Document as DocxDocument
            from docx.table import Table
            from docx.text.paragraph import Paragraph
        except ImportError as exc:
            raise DocumentParseError("缺少 python-docx 依赖，无法解析 docx") from exc

        try:
            document = DocxDocument(str(file_path))
        except Exception as exc:
            raise DocumentParseError("docx 文档解析失败") from exc

        blocks: list[ParsedBlock] = []
        current_section: str | None = None

        for child in document.element.body.iterchildren():
            if child.tag.endswith("}p"):
                paragraph = Paragraph(child, document)
                text = _normalize_text(paragraph.text)
                if not text:
                    continue

                style_name = (paragraph.style.name if paragraph.style is not None else "") or ""
                heading_level = _docx_heading_level(style_name)
                if heading_level is not None:
                    current_section = text
                    blocks.append(
                        ParsedBlock(
                            text=text,
                            block_type="heading",
                            order_index=len(blocks) + 1,
                            section_title=current_section,
                            level=heading_level,
                        )
                    )
                else:
                    blocks.append(
                        ParsedBlock(
                            text=text,
                            block_type="paragraph",
                            order_index=len(blocks) + 1,
                            section_title=current_section,
                        )
                    )
                continue

            if child.tag.endswith("}tbl"):
                table = Table(child, document)
                table_text = _docx_table_to_text(table)
                if table_text:
                    blocks.append(
                        ParsedBlock(
                            text=table_text,
                            block_type="paragraph",
                            order_index=len(blocks) + 1,
                            section_title=current_section,
                        )
                    )
        return LoaderResult(blocks=blocks, loader_name=self.name)


class LegacyDocDocumentLoader:
    name = "libreoffice-doc"

    def supports(self, file_path: Path) -> bool:
        return file_path.suffix.lower() == ".doc"

    def load(self, file_path: Path) -> LoaderResult:
        from backend.app.services.office_conversion import (
            OfficeConversionError,
            cleanup_conversion_output,
            convert_with_libreoffice_detailed,
        )

        try:
            conversion = convert_with_libreoffice_detailed(file_path, "docx")
        except OfficeConversionError as exc:
            raise DocumentParseError(str(exc)) from exc

        try:
            result = DocxDocumentLoader().load(conversion.path)
        finally:
            cleanup_conversion_output(conversion.path)
        return LoaderResult(
            blocks=result.blocks,
            loader_name=self.name,
            metadata={
                "converted_from": "doc",
                "conversion_loader": result.loader_name,
                "parser_backend": "libreoffice",
                "degraded": False,
                "degradation_reason": None,
                "conversion_elapsed_ms": conversion.elapsed_ms,
            },
        )


class AntiwordDocDocumentLoader:
    name = "antiword-doc"

    def supports(self, file_path: Path) -> bool:
        return file_path.suffix.lower() == ".doc"

    def load(self, file_path: Path) -> LoaderResult:
        import shutil
        import subprocess
        from time import perf_counter

        executable = shutil.which("antiword")
        if not executable:
            raise DocumentParseError("antiword executable not found")
        started = perf_counter()
        try:
            completed = subprocess.run(
                [executable, "-m", "UTF-8.txt", str(file_path)],
                check=False,
                capture_output=True,
                timeout=60,
            )
        except (OSError, subprocess.TimeoutExpired) as exc:
            raise DocumentParseError(f"antiword extraction failed: {exc}") from exc
        if completed.returncode != 0:
            detail = completed.stderr.decode("utf-8", errors="replace").strip()
            raise DocumentParseError(f"antiword extraction failed: {detail or completed.returncode}")
        text = completed.stdout.decode("utf-8", errors="replace")
        normalized = _normalize_text(text)
        if not normalized:
            raise DocumentParseError("antiword did not produce text")
        blocks = [
            ParsedBlock(text=paragraph, block_type="paragraph", order_index=index)
            for index, paragraph in enumerate(normalized.split("\n\n"), start=1)
            if paragraph.strip()
        ]
        return LoaderResult(
            blocks=blocks,
            loader_name=self.name,
            metadata={
                "parser_backend": "antiword",
                "degraded": True,
                "degradation_reason": "LibreOffice conversion unavailable or failed; text-only extraction used",
                "conversion_elapsed_ms": round((perf_counter() - started) * 1000, 2),
            },
        )


class PyMuPDF4LLMLoader:
    name = "pymupdf4llm"

    def supports(self, file_path: Path) -> bool:
        return file_path.suffix.lower() == ".pdf"

    def load(self, file_path: Path) -> LoaderResult:
        try:
            import pymupdf4llm
        except ImportError as exc:
            raise DocumentParseError("缺少 pymupdf4llm 依赖，无法使用版式感知 PDF loader") from exc

        try:
            pages = pymupdf4llm.to_markdown(
                str(file_path),
                page_chunks=True,
                use_ocr=False,
                show_progress=False,
                write_images=False,
                embed_images=False,
            )
        except Exception as exc:
            raise DocumentParseError("PyMuPDF4LLM 解析失败") from exc

        blocks: list[ParsedBlock] = []
        if isinstance(pages, str):
            blocks.extend(_markdown_to_blocks(pages))
        else:
            for index, page in enumerate(pages, start=1):
                metadata = page.get("metadata") or {}
                page_number = metadata.get("page_number") or index
                text = page.get("text") or ""
                page_blocks = _markdown_to_blocks(text, page_number=int(page_number))
                blocks.extend(page_blocks)
        return LoaderResult(blocks=blocks, loader_name=self.name)


class DoclingDocumentLoader:
    name = "docling"

    def supports(self, file_path: Path) -> bool:
        return file_path.suffix.lower() in {".pdf", ".docx"}

    def load(self, file_path: Path) -> LoaderResult:
        try:
            from docling.document_converter import DocumentConverter
        except ImportError as exc:
            raise DocumentParseError("缺少 docling 依赖，无法使用 Docling loader") from exc

        try:
            result = DocumentConverter().convert(str(file_path))
            markdown = result.document.export_to_markdown()
        except Exception as exc:
            raise DocumentParseError("Docling 解析失败") from exc

        return LoaderResult(blocks=_markdown_to_blocks(markdown), loader_name=self.name)


class UnstructuredDocumentLoader:
    name = "unstructured"

    def supports(self, file_path: Path) -> bool:
        return file_path.suffix.lower() in {".txt", ".md", ".docx", ".pdf"}

    def load(self, file_path: Path) -> LoaderResult:
        try:
            from unstructured.partition.auto import partition
        except ImportError as exc:
            raise DocumentParseError("缺少 unstructured 依赖，无法使用 Unstructured loader") from exc

        try:
            elements = partition(filename=str(file_path), strategy="fast")
        except Exception as exc:
            raise DocumentParseError("Unstructured 解析失败") from exc

        blocks: list[ParsedBlock] = []
        current_section: str | None = None
        for element in elements:
            text = _normalize_text(str(element))
            if not text:
                continue

            element_type = getattr(element, "category", None) or element.__class__.__name__
            metadata = getattr(element, "metadata", None)
            page_number = getattr(metadata, "page_number", None) if metadata is not None else None
            block_type = _unstructured_block_type(element_type)
            if block_type is None:
                continue
            if block_type == "heading":
                current_section = text

            blocks.append(
                ParsedBlock(
                    text=text,
                    block_type=block_type,
                    order_index=len(blocks) + 1,
                    page_number=page_number,
                    section_title=text if block_type == "heading" else current_section,
                    level=1 if block_type == "heading" else None,
                )
            )
        return LoaderResult(blocks=blocks, loader_name=self.name)


class PypdfDocumentLoader:
    name = "pypdf"

    def supports(self, file_path: Path) -> bool:
        return file_path.suffix.lower() == ".pdf"

    def load(self, file_path: Path) -> LoaderResult:
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise DocumentParseError("缺少 pypdf 依赖，无法解析 pdf") from exc

        try:
            reader = PdfReader(str(file_path))
        except Exception as exc:
            raise DocumentParseError("pdf 文档解析失败") from exc

        blocks: list[ParsedBlock] = []
        for page_number, page in enumerate(reader.pages, start=1):
            text = _normalize_text(page.extract_text() or "")
            if text:
                blocks.append(
                    ParsedBlock(
                        text=text,
                        block_type="page",
                        order_index=len(blocks) + 1,
                        page_number=page_number,
                    )
                )
        return LoaderResult(blocks=blocks, loader_name=self.name)


def _candidate_loaders(file_path: Path, loader_name: str | None = None) -> list[DocumentLoader]:
    suffix = file_path.suffix.lower()
    registry: dict[str, DocumentLoader] = {
        "text": TextDocumentLoader(),
        "markdown": MarkdownDocumentLoader(),
        "spreadsheet-csv": CsvDocumentLoader(),
        "jsonl": JsonlDocumentLoader(),
        "html": HtmlDocumentLoader(),
        "libreoffice-doc": LegacyDocDocumentLoader(),
        "antiword-doc": AntiwordDocDocumentLoader(),
        "python-docx": DocxDocumentLoader(),
        "pymupdf4llm": PyMuPDF4LLMLoader(),
        "docling": DoclingDocumentLoader(),
        "unstructured": UnstructuredDocumentLoader(),
        "pypdf": PypdfDocumentLoader(),
    }
    loaders = [
        registry[name]
        for name in DOCUMENT_LOADER_ORDER.get(suffix, [])
        if name in registry and registry[name].supports(file_path)
    ]

    if loader_name is None:
        return loaders
    return [loader for loader in loaders if loader.name == loader_name]


class _StructuredHtmlParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.blocks: list[ParsedBlock] = []
        self.title: str | None = None
        self.current_section: str | None = None
        self._capture_tag: str | None = None
        self._capture: list[str] = []
        self._title_capture: list[str] | None = None
        self._skip_depth = 0
        self._row: list[str] | None = None
        self._cell: list[str] | None = None

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        tag = tag.lower()
        if tag in {"script", "style", "noscript"}:
            self._skip_depth += 1
            return
        if self._skip_depth:
            return
        if tag == "title":
            self._title_capture = []
        elif tag == "tr":
            self._row = []
        elif tag in {"td", "th"} and self._row is not None:
            self._cell = []
        elif tag in {"h1", "h2", "h3", "h4", "h5", "h6", "p", "li"} and self._cell is None:
            self._capture_tag = tag
            self._capture = []

    def handle_endtag(self, tag: str) -> None:
        tag = tag.lower()
        if tag in {"script", "style", "noscript"}:
            self._skip_depth = max(0, self._skip_depth - 1)
            return
        if self._skip_depth:
            return
        if tag == "title":
            self.title = _normalize_text(" ".join(self._title_capture or [])) or None
            self._title_capture = None
        elif tag in {"td", "th"} and self._cell is not None and self._row is not None:
            cell_text = _normalize_text(" ".join(self._cell))
            if cell_text:
                self._row.append(cell_text)
            self._cell = None
        elif tag == "tr" and self._row is not None:
            if self._row:
                self.blocks.append(
                    ParsedBlock(
                        text="；".join(self._row),
                        block_type="paragraph",
                        order_index=len(self.blocks) + 1,
                        section_title=self.current_section,
                    )
                )
            self._row = None
        elif tag == self._capture_tag:
            text = _normalize_text(" ".join(self._capture))
            if text:
                is_heading = tag.startswith("h")
                if is_heading:
                    self.current_section = text
                self.blocks.append(
                    ParsedBlock(
                        text=text,
                        block_type="heading" if is_heading else "paragraph",
                        order_index=len(self.blocks) + 1,
                        section_title=self.current_section,
                        level=int(tag[1]) if is_heading else None,
                    )
                )
            self._capture_tag = None
            self._capture = []

    def handle_data(self, data: str) -> None:
        if self._skip_depth:
            return
        if self._cell is not None:
            self._cell.append(data)
        elif self._capture_tag:
            self._capture.append(data)
        elif self._title_capture is not None:
            self._title_capture.append(data)


def _json_metadata_value(value: Any) -> bool:
    if value is None:
        return False
    if isinstance(value, (str, int, float, bool)):
        return True
    if isinstance(value, list):
        return len(value) <= 100
    if isinstance(value, dict):
        return len(value) <= 50
    return False


def _read_utf8_text(file_path: Path) -> str:
    try:
        return file_path.read_text(encoding="utf-8-sig")
    except UnicodeDecodeError as exc:
        raise DocumentParseError("文档编码不是 UTF-8") from exc


def _markdown_to_blocks(text: str, page_number: int | None = None) -> list[ParsedBlock]:
    blocks: list[ParsedBlock] = []
    current_section: str | None = None
    paragraph_lines: list[str] = []

    def flush_paragraph() -> None:
        nonlocal paragraph_lines
        paragraph = _normalize_text("\n".join(paragraph_lines))
        if paragraph:
            blocks.append(
                ParsedBlock(
                    text=paragraph,
                    block_type="paragraph",
                    order_index=len(blocks) + 1,
                    page_number=page_number,
                    section_title=current_section,
                )
            )
        paragraph_lines = []

    for line in text.splitlines():
        stripped = line.strip()
        heading_level = _markdown_heading_level(stripped)
        if heading_level is not None:
            flush_paragraph()
            title = stripped.lstrip("#").strip()
            current_section = title
            blocks.append(
                ParsedBlock(
                    text=title,
                    block_type="heading",
                    order_index=len(blocks) + 1,
                    page_number=page_number,
                    section_title=title,
                    level=heading_level,
                )
            )
            continue

        if not stripped:
            flush_paragraph()
            continue

        paragraph_lines.append(line)

    flush_paragraph()
    return blocks


def _markdown_heading_level(stripped_line: str) -> int | None:
    if not stripped_line.startswith("#") or len(stripped_line) > 120:
        return None
    marker = stripped_line.split(" ", maxsplit=1)[0]
    if not marker or any(char != "#" for char in marker) or len(marker) > 6:
        return None
    if len(stripped_line) == len(marker):
        return None
    return len(marker)


def _unstructured_block_type(element_type: str) -> str | None:
    if element_type in {"Title", "Header"}:
        return "heading"
    if element_type == "Table":
        return "table"
    if element_type in {"Footer", "PageBreak"}:
        return None
    return "paragraph"


def _docx_heading_level(style_name: str) -> int | None:
    normalized = style_name.lower()
    if normalized.startswith("heading"):
        parts = normalized.split()
        if len(parts) > 1 and parts[-1].isdigit():
            return int(parts[-1])
        return 1
    return None


def _docx_table_to_text(table) -> str:
    rows: list[str] = []
    for row in table.rows:
        cells: list[str] = []
        for cell in row.cells:
            cell_text = _normalize_text("\n".join(paragraph.text for paragraph in cell.paragraphs))
            if cell_text:
                cells.append(cell_text)
        if cells:
            rows.append("；".join(cells))
    return "\n".join(rows).strip()


def _normalize_text(text: str) -> str:
    lines = [line.strip() for line in text.replace("\r\n", "\n").replace("\r", "\n").split("\n")]
    compact_lines: list[str] = []
    previous_blank = False
    for line in lines:
        if not line:
            if not previous_blank:
                compact_lines.append("")
            previous_blank = True
            continue
        compact_lines.append(" ".join(line.split()))
        previous_blank = False
    return "\n".join(compact_lines).strip()


def _ensure_blocks(blocks: list[ParsedBlock]) -> list[ParsedBlock]:
    cleaned = [block for block in blocks if block.text.strip()]
    if not cleaned:
        raise DocumentParseError("文档没有可解析文本")
    for index, block in enumerate(cleaned, start=1):
        block.order_index = index
    return cleaned


def _join_blocks(blocks: list[ParsedBlock]) -> str:
    return "\n\n".join(block.text for block in blocks if block.text.strip()).strip()
