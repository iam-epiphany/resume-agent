import pytest

from backend.app.models.document import Document, DocumentChunk
from backend.app.services.chunk_service import build_chunks_from_parsed, count_tokens
from backend.app.services.document_types import ParsedBlock, ParsedDocument
from backend.app.services.document_parser import DocumentParseError, available_loader_names, parse_document
from backend.app.services.loader_evaluation import evaluate_document_loaders


def test_parse_markdown_returns_structured_blocks(tmp_path) -> None:
    path = tmp_path / "rules.md"
    path.write_text("# 材料填写说明\n\n## 资产合计\n资产合计应等于各项资产分项金额合计。\n", encoding="utf-8")

    parsed = parse_document(path)

    assert parsed.metadata["parser_version"] == "structured-v4-formula"
    assert parsed.metadata["loader_name"] == "markdown"
    assert [block.block_type for block in parsed.blocks] == ["heading", "heading", "paragraph"]
    assert parsed.blocks[2].section_title == "资产合计"


def test_parse_jsonl_requires_content_and_rejects_qa_answers(tmp_path) -> None:
    valid = tmp_path / "rules.jsonl"
    valid.write_text('{"title":"规则一","text":"技能掌握程度应符合材料要求。"}\n', encoding="utf-8")
    parsed = parse_document(valid)
    assert parsed.blocks[0].section_title == "规则一"

    qa = tmp_path / "qa.jsonl"
    qa.write_text('{"question":"答案是什么","answer":"A","evidence":"标准答案"}\n', encoding="utf-8")
    with pytest.raises(DocumentParseError, match="QA/答案数据"):
        parse_document(qa)


def test_parse_html_preserves_headings_paragraphs_and_tables_as_plain_text(tmp_path) -> None:
    path = tmp_path / "rule.html"
    path.write_text(
        "<html><head><title>材料规则</title></head><body><h1>第一章</h1><p>正文依据。</p>"
        "<table><tr><th>指标</th><th>值</th></tr><tr><td>资本</td><td>12</td></tr></table></body></html>",
        encoding="utf-8",
    )

    parsed = parse_document(path)

    assert parsed.metadata["html_title"] == "材料规则"
    assert any(block.block_type == "heading" and block.text == "第一章" for block in parsed.blocks)
    assert any(block.block_type == "paragraph" and "资本" in block.text for block in parsed.blocks)


def test_loader_order_is_configured_by_file_type(tmp_path) -> None:
    assert available_loader_names(tmp_path / "rules.txt") == ["text", "unstructured"]
    assert available_loader_names(tmp_path / "rules.md") == ["markdown", "unstructured"]
    assert available_loader_names(tmp_path / "rules.doc") == ["libreoffice-doc", "antiword-doc"]
    assert available_loader_names(tmp_path / "rules.docx") == ["python-docx", "docling", "unstructured"]
    assert available_loader_names(tmp_path / "rules.pdf") == ["pymupdf4llm", "docling", "unstructured", "pypdf"]


def test_chunks_inherit_section_from_structured_blocks(tmp_path) -> None:
    path = tmp_path / "rules.md"
    path.write_text("# 材料填写说明\n\n## 资产合计\n资产合计应等于各项资产分项金额合计。\n", encoding="utf-8")
    parsed = parse_document(path)

    chunks = build_chunks_from_parsed(document_id="DOC-TEST-0001", parsed=parsed)

    assert len(chunks) == 1
    assert chunks[0].section_title == "资产合计"
    assert "资产合计应等于各项资产分项金额合计" in chunks[0].text
    assert "章节路径：材料填写说明 > 资产合计" in chunks[0].embedding_text
    assert "章节：资产合计" in chunks[0].embedding_text
    assert "内容类型：正文" in chunks[0].embedding_text
    assert chunks[0].token_count > 0


def test_contextual_embedding_text_includes_document_and_location_metadata() -> None:
    parsed = ParsedDocument(
        text="3.1 资产合计差异处理\n\n若差异来自外币折算，应保留汇率日期。",
        metadata={"source_format": "pdf"},
        blocks=[
            ParsedBlock(
                text="3.1 资产合计差异处理",
                block_type="heading",
                order_index=1,
                page_number=7,
                section_title="3.1 资产合计差异处理",
                level=2,
            ),
            ParsedBlock(
                text="若差异来自外币折算，应保留汇率日期。",
                block_type="paragraph",
                order_index=2,
                page_number=7,
                section_title="3.1 资产合计差异处理",
            ),
        ],
    )

    chunks = build_chunks_from_parsed(
        document_id="DOC-TEST-0001",
        parsed=parsed,
        source_file="技能掌握情况统计表填写说明.pdf",
    )

    assert len(chunks) == 1
    assert chunks[0].text == "3.1 资产合计差异处理\n\n若差异来自外币折算，应保留汇率日期。"
    assert "来源文件：技能掌握情况统计表填写说明.pdf" in chunks[0].embedding_text
    assert "文档格式：pdf" in chunks[0].embedding_text
    assert "章节路径：3.1 资产合计差异处理" in chunks[0].embedding_text
    assert "条款号：3.1" in chunks[0].embedding_text
    assert "父条款号：3" in chunks[0].embedding_text
    assert "页码：7" in chunks[0].embedding_text
    assert chunks[0].embedding_text.endswith(chunks[0].text)


def test_chunks_keep_section_hierarchy_and_neighbor_ids(tmp_path) -> None:
    path = tmp_path / "rules.md"
    path.write_text(
        "## 3. 资产合计与校验关系\n"
        "资产合计应等于各项资产分项金额之和。\n\n"
        "### 3.1 资产合计差异处理\n"
        "若差异来自外币折算，应保留汇率日期、折算规则和原币金额来源。\n",
        encoding="utf-8",
    )
    parsed = parse_document(path)

    chunks = build_chunks_from_parsed(document_id="DOC-TEST-0001", parsed=parsed)

    assert len(chunks) == 2
    assert chunks[0].section_number == "3"
    assert chunks[0].parent_section_number is None
    assert chunks[0].section_path == ["3. 资产合计与校验关系"]
    assert chunks[0].next_chunk_id == chunks[1].chunk_id
    assert chunks[1].section_number == "3.1"
    assert chunks[1].parent_section_number == "3"
    assert chunks[1].section_path == ["3. 资产合计与校验关系", "3.1 资产合计差异处理"]
    assert chunks[1].previous_chunk_id == chunks[0].chunk_id


def test_chunk_splitting_respects_max_tokens_and_overlap() -> None:
    text = "资产质量" * 900
    parsed = ParsedDocument(
        text=text,
        blocks=[ParsedBlock(text=text, block_type="paragraph", order_index=1, section_title="资产质量")],
    )

    chunks = build_chunks_from_parsed(document_id="DOC-TEST-0001", parsed=parsed)

    assert len(chunks) > 1
    assert all(chunk.token_count <= 800 for chunk in chunks)
    assert chunks[0].text[-10:] in chunks[1].text


def test_long_paragraph_prefers_sentence_boundaries() -> None:
    sentence_a = "资产合计应等于各项资产分项金额合计。"
    sentence_b = "填写人员应保留表述说明和数据来源。"
    text = (sentence_a + sentence_b) * 90
    parsed = ParsedDocument(
        text=text,
        blocks=[ParsedBlock(text=text, block_type="paragraph", order_index=1, section_title="资产合计")],
    )

    chunks = build_chunks_from_parsed(document_id="DOC-TEST-0001", parsed=parsed)

    assert len(chunks) > 1
    assert all(count_tokens(chunk.text) <= 800 for chunk in chunks)
    assert all(chunk.text.endswith("。") for chunk in chunks)
    assert not any(chunk.text.startswith("合计。") for chunk in chunks[1:])


def test_semantic_break_splits_topic_shift(monkeypatch) -> None:
    paragraphs = [
        "资产合计应等于各项资产分项金额合计。" * 80,
        "技能分类应按照材料规定保持表述一致。" * 80,
        "成绩排名指标应按照材料口径计算。" * 80,
    ]
    parsed = ParsedDocument(
        text="\n\n".join(paragraphs),
        blocks=[
            ParsedBlock(
                text=paragraph,
                block_type="paragraph",
                order_index=index,
                section_title="材料指标",
            )
            for index, paragraph in enumerate(paragraphs, start=1)
        ],
    )

    monkeypatch.setattr(
        "backend.app.services.chunk_service.embed_for_semantic_split",
        lambda texts: [[1.0, 0.0], [0.95, 0.05], [0.0, 1.0]],
    )

    chunks = build_chunks_from_parsed(document_id="DOC-TEST-0001", parsed=parsed)

    assert len(chunks) >= 2
    assert any("成绩排名" in chunk.text for chunk in chunks)
    assert all(count_tokens(chunk.text) <= 800 for chunk in chunks)


def test_semantic_chunking_falls_back_when_embedding_is_unavailable(monkeypatch) -> None:
    from backend.app.services.embedding_service import EmbeddingServiceError

    paragraphs = [
        "资产合计应等于各项资产分项金额合计。" * 80,
        "成绩排名指标应按照材料口径计算。" * 80,
    ]
    parsed = ParsedDocument(
        text="\n\n".join(paragraphs),
        blocks=[
            ParsedBlock(
                text=paragraph,
                block_type="paragraph",
                order_index=index,
                section_title="材料指标",
            )
            for index, paragraph in enumerate(paragraphs, start=1)
        ],
    )
    monkeypatch.setattr(
        "backend.app.services.chunk_service.embed_for_semantic_split",
        lambda texts: (_ for _ in ()).throw(EmbeddingServiceError("model unavailable")),
    )

    chunks = build_chunks_from_parsed(document_id="DOC-TEST-0001", parsed=parsed)

    assert chunks
    assert any("成绩排名" in chunk.text for chunk in chunks)
    assert all(count_tokens(chunk.text) <= 800 for chunk in chunks)


def test_docx_loader_preserves_paragraph_and_table_order_as_plain_text(tmp_path) -> None:
    from docx import Document as DocxDocument

    path = tmp_path / "rules.docx"
    document = DocxDocument()
    document.add_paragraph("一、填写说明")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "指标"
    table.cell(0, 1).text = "口径"
    table.cell(1, 0).text = "资产合计"
    table.cell(1, 1).text = "资产分项合计"
    document.add_paragraph("二、留痕要求")
    document.save(path)

    parsed = parse_document(path)

    assert [block.block_type for block in parsed.blocks] == ["paragraph", "paragraph", "paragraph"]
    assert parsed.blocks[0].text == "一、填写说明"
    assert "资产合计" in parsed.blocks[1].text
    assert "资产分项合计" in parsed.blocks[1].text
    assert parsed.blocks[2].text == "二、留痕要求"


def test_office_conversion_rejects_oversized_output(tmp_path, monkeypatch) -> None:
    from backend.app.services import office_conversion

    output_dir = tmp_path / "resumemind-office-test"
    output_dir.mkdir()
    converted = output_dir / "oversized.docx"
    converted.write_bytes(b"123456789")
    monkeypatch.setattr(office_conversion, "OFFICE_CONVERSION_MAX_BYTES", 8)

    with pytest.raises(office_conversion.OfficeConversionError, match="exceeds 8 bytes"):
        office_conversion._validate_conversion_output(converted, output_dir)

    assert not output_dir.exists()


def test_legacy_doc_loader_uses_libreoffice_conversion(tmp_path, monkeypatch) -> None:
    from docx import Document as DocxDocument

    converted_path = tmp_path / "converted.docx"
    document = DocxDocument()
    document.add_paragraph("简历材料正文")
    document.save(converted_path)
    legacy_path = tmp_path / "legacy.doc"
    legacy_path.write_bytes(b"legacy-binary-placeholder")

    from backend.app.services.office_conversion import OfficeConversionResult

    monkeypatch.setattr(
        "backend.app.services.office_conversion.convert_with_libreoffice_detailed",
        lambda file_path, target_extension: OfficeConversionResult(
            path=converted_path, backend="libreoffice", elapsed_ms=12.5
        ),
    )

    parsed = parse_document(legacy_path)

    assert parsed.metadata["loader_name"] == "libreoffice-doc"
    assert parsed.metadata["converted_from"] == "doc"
    assert parsed.metadata["parser_backend"] == "libreoffice"
    assert parsed.metadata["degraded"] is False
    assert "简历材料正文" in parsed.text


def test_legacy_doc_falls_back_to_antiword(tmp_path, monkeypatch) -> None:
    import subprocess

    legacy_path = tmp_path / "legacy.doc"
    legacy_path.write_bytes(b"legacy-binary-placeholder")

    monkeypatch.setattr(
        "backend.app.services.office_conversion.convert_with_libreoffice_detailed",
        lambda file_path, target_extension: (_ for _ in ()).throw(
            __import__(
                "backend.app.services.office_conversion", fromlist=["OfficeConversionError"]
            ).OfficeConversionError("forced failure")
        ),
    )
    monkeypatch.setattr("shutil.which", lambda name: "antiword" if name == "antiword" else None)
    monkeypatch.setattr(
        subprocess,
        "run",
        lambda *args, **kwargs: subprocess.CompletedProcess(
            args=args[0], returncode=0, stdout="简历材料降级正文".encode("utf-8"), stderr=b""
        ),
    )

    parsed = parse_document(legacy_path)

    assert parsed.metadata["loader_name"] == "antiword-doc"
    assert parsed.metadata["parser_backend"] == "antiword"
    assert parsed.metadata["degraded"] is True
    assert "简历材料降级正文" in parsed.text


def test_pdf_defaults_to_pymupdf4llm_loader(tmp_path) -> None:
    import pymupdf

    path = tmp_path / "rules.pdf"
    document = pymupdf.open()
    page = document.new_page()
    page.insert_text((72, 72), "Total assets equal the sum of asset items.")
    document.save(path)
    document.close()

    parsed = parse_document(path)

    assert parsed.metadata["loader_name"] in {"pymupdf4llm", "pypdf"}
    assert parsed.blocks[0].page_number == 1
    assert "Total assets" in parsed.text


def test_loader_evaluation_compares_pdf_loaders(tmp_path) -> None:
    import pymupdf

    path = tmp_path / "rules.pdf"
    document = pymupdf.open()
    page = document.new_page()
    page.insert_text((72, 72), "材料填写说明")
    document.save(path)
    document.close()

    evaluations = evaluate_document_loaders(path)

    names = [evaluation.loader_name for evaluation in evaluations]
    assert names == ["pymupdf4llm", "docling", "unstructured", "pypdf"]
    assert any(evaluation.ok and evaluation.block_count > 0 for evaluation in evaluations)


